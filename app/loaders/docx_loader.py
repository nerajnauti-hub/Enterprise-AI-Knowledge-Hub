from pathlib import Path
from docx import Document as WordDocument

from app.models.document import Document


class DOCXLoader:

    def load(self, filepath):

        doc = WordDocument(filepath)

        paragraphs = []

        for paragraph in doc.paragraphs:

            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        content = "\n\n".join(paragraphs)

        return Document(
            filename=Path(filepath).name,
            filetype=".docx",
            content=content,
            metadata={}
        )